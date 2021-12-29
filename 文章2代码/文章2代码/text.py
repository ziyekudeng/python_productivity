from pathlib import Path
import docx


def merge_without_format(docx_files):
    '''
    只获取内容进行合并
    '''
    # 遍历每个文件
    for docx_file in sorted(docx_files):
        another_doc = Document(docx_files)
        # 获取每个文件的所有"段落"
        paras = another_doc.paragraphs
        # 获取所有段落的文字内容
        for para in paras:
            # 为新的word创建一个新段落
            newpar = doc.add_paragraph('')
            # 将提取的内容写入新的文本段落中
            newpar.add_run(para.text)

    # 所有文件合并完成后在指定路径进行保存
    doc.save(Path(word_files_path,'new.docx'))

# 调用函数
merge_without_format(files)